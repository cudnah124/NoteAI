"""Document processing logic
Handles text extraction, chunking, and metadata
"""
import hashlib
from typing import List, Dict, Any
from pypdf import PdfReader
from io import BytesIO
import logging
import re
from urllib.parse import urlparse, parse_qs
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document types"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def calculate_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of content"""
        return hashlib.sha256(content).hexdigest()
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Chunk text into smaller pieces with overlap
        
        Args:
            text: Input text
            metadata: Optional metadata to attach to chunks
            
        Returns:
            List of chunk dictionaries with content and metadata
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            chunk_meta = metadata.copy() if metadata else {}
            chunk_meta["chunk_index"] = len(chunks)
            chunk_meta["start_word"] = i
            chunk_meta["end_word"] = i + len(chunk_words)
            
            chunks.append({
                "content": chunk_text,
                "metadata": chunk_meta
            })
        
        return chunks
    
    def process_pdf(self, file_content: bytes) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process PDF file
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Tuple of (full_text, chunks)
        """
        pdf_file = BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        all_chunks = []
        full_text = ""
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            full_text += page_text + "\n"
            
            # Chunk this page
            page_chunks = self.chunk_text(
                page_text,
                metadata={"page_num": page_num + 1}
            )
            all_chunks.extend(page_chunks)
        
        return full_text, all_chunks
    
    def process_docx(self, file_content: bytes) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process DOCX file
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Tuple of (full_text, chunks)
        """
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            all_chunks = []
            full_text = ""
            
            # Extract text from paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if para_text.strip():  # Skip empty paragraphs
                    full_text += para_text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        full_text += row_text + "\n"
            
            # Chunk the entire document
            all_chunks = self.chunk_text(
                full_text,
                metadata={"type": "docx"}
            )
            
            logger.info(f"DOCX processed successfully, extracted {len(full_text)} characters, {len(all_chunks)} chunks")
            
            return full_text, all_chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise
    
    def process_image(self, file_content: bytes) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process image file (placeholder - needs OCR integration)
        
        Args:
            file_content: Image file bytes
            
        Returns:
            Tuple of (extracted_text, chunks)
        """
        # TODO: Integrate Clova OCR
        text = "[Image OCR not yet implemented]"
        chunks = self.chunk_text(text, metadata={"type": "image"})
        return text, chunks
    
    def process_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Process plain text
        
        Args:
            text: Input text
            metadata: Optional metadata
            
        Returns:
            List of chunks
        """
        return self.chunk_text(text, metadata)
    
    def extract_youtube_video_id(self, url: str) -> str:
        """
        Extract YouTube video ID from URL
        
        Args:
            url: YouTube URL
            
        Returns:
            Video ID string
        """
        # Handle youtu.be short URLs
        if 'youtu.be/' in url:
            return url.split('youtu.be/')[-1].split('?')[0]
        
        # Handle youtube.com URLs
        parsed = urlparse(url)
        if parsed.hostname in ['www.youtube.com', 'youtube.com', 'm.youtube.com']:
            if parsed.path == '/watch':
                query_params = parse_qs(parsed.query)
                return query_params.get('v', [None])[0]
            elif parsed.path.startswith('/embed/'):
                return parsed.path.split('/embed/')[-1].split('?')[0]
            elif parsed.path.startswith('/v/'):
                return parsed.path.split('/v/')[-1].split('?')[0]
        
        return None
    
    async def process_youtube_url(self, url: str) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process YouTube URL using YouTube Transcript API (fast method)
        
        Args:
            url: YouTube video URL
            
        Returns:
            Tuple of (transcript_text, chunks)
        """
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
            
            # Extract video ID
            video_id = self.extract_youtube_video_id(url)
            if not video_id:
                logger.error(f"Could not extract video ID from URL: {url}")
                text = "[Invalid YouTube URL]"
                chunks = self.chunk_text(text, metadata={"type": "youtube", "error": "invalid_url"})
                return text, chunks
            
            logger.info(f"Fetching transcript for YouTube video: {video_id}")
            
            # Try to get transcript (prioritize Vietnamese, then English, then any available)
            try:
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                
                # Try Vietnamese first
                try:
                    transcript = transcript_list.find_transcript(['vi'])
                except:
                    # Try English
                    try:
                        transcript = transcript_list.find_transcript(['en'])
                    except:
                        # Get any available transcript
                        transcript = transcript_list.find_generated_transcript(['vi', 'en', 'ko'])
                
                # Fetch the actual transcript
                transcript_data = transcript.fetch()
                
            except (TranscriptsDisabled, NoTranscriptFound):
                logger.warning(f"No transcript available for video {video_id}, trying direct fetch")
                # Fallback: try direct fetch with language codes
                try:
                    transcript_data = YouTubeTranscriptApi.get_transcript(video_id, languages=['vi', 'en', 'ko'])
                except:
                    logger.error(f"All transcript methods failed for video {video_id}")
                    text = "[YouTube transcript not available for this video]"
                    chunks = self.chunk_text(text, metadata={"type": "youtube", "error": "no_transcript"})
                    return text, chunks
            
            # Combine transcript segments
            full_text = " ".join([entry['text'] for entry in transcript_data])
            
            if not full_text.strip():
                logger.error("Transcript is empty")
                text = "[YouTube transcript is empty]"
                chunks = self.chunk_text(text, metadata={"type": "youtube", "error": "empty_transcript"})
                return text, chunks
            
            logger.info(f"YouTube transcript fetched successfully, length: {len(full_text)} characters")
            
            # Chunk the transcript with timestamps
            chunks = self.chunk_text(
                full_text, 
                metadata={
                    "type": "youtube", 
                    "source": "youtube_transcript_api",
                    "video_id": video_id
                }
            )
            
            return full_text, chunks
            
        except Exception as e:
            logger.error(f"Error processing YouTube URL: {e}")
            text = f"[YouTube processing error: {str(e)}]"
            chunks = self.chunk_text(text, metadata={"type": "youtube", "error": str(e)})
            return text, chunks
    
    async def process_web_url(self, url: str) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process web URL using BeautifulSoup for fast HTML parsing
        
        Args:
            url: Website URL
            
        Returns:
            Tuple of (extracted_text, chunks)
        """
        try:
            import httpx
            from bs4 import BeautifulSoup
            
            logger.info(f"Fetching content from URL: {url}")
            
            # Fetch URL with timeout and proper headers to avoid 403 errors
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }
            
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True, headers=headers) as client:
                response = await client.get(url)
                response.raise_for_status()
                html_content = response.text
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text(separator='\n', strip=True)
            
            # Clean up text: remove multiple newlines and spaces
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            full_text = '\n'.join(lines)
            
            if not full_text.strip():
                logger.error("No text content extracted from URL")
                text = "[No text content found on webpage]"
                chunks = self.chunk_text(text, metadata={"type": "web", "error": "empty_content"})
                return text, chunks
            
            logger.info(f"Web content extracted successfully, length: {len(full_text)} characters")
            
            # Chunk the content
            chunks = self.chunk_text(
                full_text,
                metadata={
                    "type": "web",
                    "source": "web_scraping",
                    "url": url
                }
            )
            
            return full_text, chunks
            
        except httpx.TimeoutException:
            logger.error(f"Timeout fetching URL: {url}")
            text = "[Website request timed out]"
            chunks = self.chunk_text(text, metadata={"type": "web", "error": "timeout"})
            return text, chunks
        except httpx.HTTPError as e:
            logger.error(f"HTTP error fetching URL: {e}")
            text = f"[HTTP error: {str(e)}]"
            chunks = self.chunk_text(text, metadata={"type": "web", "error": str(e)})
            return text, chunks
        except Exception as e:
            logger.error(f"Error processing web URL: {e}")
            text = f"[Web processing error: {str(e)}]"
            chunks = self.chunk_text(text, metadata={"type": "web", "error": str(e)})
            return text, chunks
    
    async def process_video(self, file_content: bytes) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process video file using Naver Clova Speech for transcription
        
        Args:
            file_content: Video file bytes
            
        Returns:
            Tuple of (transcript_text, chunks)
        """
        try:
            from app.integrations.naver.whisper import WhisperService
            
            whisper = WhisperService()
            logger.info("Starting video transcription with Naver Clova Speech API...")
            
            # Transcribe video to text
            transcript = await whisper.transcribe_video(file_content, language="ko-KR")
            
            if not transcript:
                logger.error("Video transcription failed")
                text = "[Video transcription failed]"
                chunks = self.chunk_text(text, metadata={"type": "video", "error": "transcription_failed"})
                return text, chunks
            
            logger.info(f"Transcription successful, length: {len(transcript)} characters")
            
            # Chunk the transcript
            chunks = self.chunk_text(transcript, metadata={"type": "video", "source": "naver_clova_speech"})
            
            return transcript, chunks
            
        except Exception as e:
            logger.error(f"Error processing video: {e}")
            text = f"[Video processing error: {str(e)}]"
            chunks = self.chunk_text(text, metadata={"type": "video", "error": str(e)})
            return text, chunks
